"""
Exporter Mixin for Report Service.
Contains PDF and DOCX generation methods.
"""

from __future__ import annotations

import io
import os


class ReportExporterMixin:
    """
    Mixin for document conversion (PDF & DOCX).
    Expects self._logger to be present.
    """

    def _prepare_watermark(self, logo_path: str, org_id: str) -> str | None:
        """
        Creates a semi-transparent rotated watermark image from a logo.
        Returns the path to the temporary watermark file.
        """
        try:
            from PIL import Image, ImageEnhance
            import tempfile

            # 1. Path Traversal & Injection Defense: Sanitize org_id 
            # (Strip any non-alphanumeric/hyphen characters to keep it safe for filenames)
            safe_org_id = "".join(c for c in org_id if c.isalnum() or c == "-")[:64]
            if not safe_org_id:
                self._logger.error(f"Invalid org_id for watermark: {org_id}")
                return None

            # 2. Guard Check for Asset Existence
            if not os.path.exists(logo_path):
                self._logger.error(f"Watermark source logo missing at: {logo_path}")
                return None
            
            # 3. Use system temp directory (Production-safe for read-only containers/serverless)
            tmp_dir = os.path.join(tempfile.gettempdir(), "kreesalis_reports_wm")
            os.makedirs(tmp_dir, exist_ok=True)
            watermark_path = os.path.join(tmp_dir, f"watermark_{safe_org_id}.png")
            
            # Optimization: Re-use existing watermark if already generated in this lifecycle
            if os.path.exists(watermark_path):
                return watermark_path
            
            img = Image.open(logo_path).convert("RGBA")
            
            # Rotate diagonal (e.g. 45 degrees)
            rotated = img.rotate(45, expand=True, resample=Image.BICUBIC)
            
            # Reduce alpha (lighten)
            alpha = rotated.split()[3]
            alpha = ImageEnhance.Brightness(alpha).enhance(0.15) # 15% opacity
            rotated.putalpha(alpha)
            
            rotated.save(watermark_path)
            return watermark_path
        except Exception as e:
            if hasattr(self, "_logger"):
                self._logger.error(f"Failed to prepare watermark: {e}")
            return logo_path # Fallback to original logo if processing fails

    async def _convert_to_pdf(self, markdown_text: str, title: str, show_watermark: bool = False, watermark_path: str | None = None) -> bytes:
        from fpdf import FPDF

        def _sanitize(t: str) -> str:
            # fpdf2 core fonts use latin-1. Emojis and complex unicode will crash the generator.
            # We strip them cleanly to ensure the PDF generation always succeeds.
            return t.encode("latin-1", errors="ignore").decode("latin-1")

        class ReportPDF(FPDF):
            def __init__(self, *args, **kwargs):
                super().__init__(*args, **kwargs)
                self.show_watermark = show_watermark
                self.watermark_path = watermark_path or os.path.join(os.path.dirname(os.path.abspath(__file__)), "kreesalis.png")

            def header(self):
                if self.show_watermark:
                    try:
                        # Draw watermark in the center. 
                        # We don't need local_context(fill_opacity) because _prepare_watermark 
                        # already baked the transparency into the generated image.
                        self.image(self.watermark_path, x=40, y=100, w=130)
                    except Exception as e:
                        print(f"Failed to add PDF watermark: {e}")

            def footer(self):
                self.set_y(-15)
                self.set_font("helvetica", "I", 8)
                self.set_text_color(128, 128, 128)
                self.cell(0, 10, txt=f"Page {self.page_no()}/{{nb}}", align="C")

        pdf = ReportPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        
        # Add Title
        pdf.set_font("helvetica", "B", 18)
        pdf.set_text_color(30, 58, 95) # #1e3a5f
        pdf.multi_cell(0, 10, txt=_sanitize(title or "Report"), align="C")
        pdf.set_line_width(0.5)
        pdf.set_draw_color(238, 238, 238) # #eee
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
        pdf.ln(10)

        # Robust manual markdown parser layout engine
        lines = markdown_text.split("\n")
        pdf.set_text_color(51, 51, 51) # #333

        in_table = False
        table_data = []

        def flush_table():
            if not table_data:
                return
            # Render fpdf2 native table
            # fpdf2 allows rendering a table using pdf.table() ContextManager
            with pdf.table(
                borders_layout="ALL",
                cell_fill_color=(248, 248, 248), # Light gray header
                cell_fill_mode="ROWS",
                text_align="LEFT",
                line_height=6,
            ) as table:
                for idx, row in enumerate(table_data):
                    table_row = table.row()
                    for cell in row:
                        if idx == 0:
                            pdf.set_font("helvetica", "B", 10)
                        else:
                            pdf.set_font("helvetica", "", 10)
                        table_row.cell(_sanitize(cell))
            pdf.ln(5)
            table_data.clear()

        for line_raw in lines:
            line = line_raw.strip()
            
            # Handle table accumulation or rendering
            if line.startswith("|") and line.endswith("|"):
                in_table = True
                if "---" in line:
                    continue # Skip separator row
                
                # Split and clean table cells
                cells = [c.strip() for c in line.split("|")[1:-1]]
                table_data.append(cells)
                continue
            else:
                if in_table:
                    flush_table()
                    in_table = False

            if not line:
                pdf.ln(3)
                continue

            # Always reset X
            pdf.set_x(pdf.l_margin)

            if line.startswith("# "):
                pdf.ln(5)
                pdf.set_font("helvetica", "B", 16)
                pdf.set_text_color(30, 58, 95)
                heading_text = line[2:].replace("**", "").strip() or " "
                pdf.multi_cell(0, 10, txt=_sanitize(heading_text), align="L")
                pdf.set_x(pdf.l_margin)
                pdf.set_draw_color(238, 238, 238)
                pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 180, pdf.get_y())
                pdf.ln(5)
                pdf.set_text_color(51, 51, 51)
            elif line.startswith("## "):
                pdf.ln(5)
                pdf.set_font("helvetica", "B", 14)
                pdf.set_text_color(44, 82, 130)
                heading_text = line[3:].replace("**", "").strip() or " "
                pdf.multi_cell(0, 10, txt=_sanitize(heading_text), align="L")
                pdf.ln(2)
                pdf.set_text_color(51, 51, 51)
            elif line.startswith("### "):
                pdf.ln(3)
                pdf.set_font("helvetica", "B", 12)
                heading_text = line[4:].replace("**", "").strip() or " "
                pdf.multi_cell(0, 8, txt=_sanitize(heading_text), align="L")
                pdf.ln(2)
            elif line.startswith("> "):
                pdf.set_font("helvetica", "I", 10)
                pdf.set_text_color(100, 100, 100)
                text = line[2:].replace("**", "").strip()
                pdf.set_x(pdf.l_margin + 5)
                # Draw a gentle left border for the blockquote
                pdf.set_draw_color(200, 200, 200)
                pdf.set_line_width(0.8)
                y_start = pdf.get_y()
                pdf.multi_cell(0, 6, txt=_sanitize(text), align="L")
                y_end = pdf.get_y()
                pdf.line(pdf.l_margin + 2, y_start + 1, pdf.l_margin + 2, y_end - 1)
                pdf.set_text_color(51, 51, 51)
            elif line.startswith("- ") or line.startswith("* "):
                pdf.set_font("helvetica", "", 10)
                text = line[2:].replace("**", "").strip()
                bullet_str = f"•  {text}" if text else "•"
                sanitized_bullet = _sanitize(bullet_str)
                if not sanitized_bullet.strip():
                    sanitized_bullet = " "
                pdf.multi_cell(0, 6, txt=sanitized_bullet, align="L")
            else:
                pdf.set_font("helvetica", "", 10)
                text = line.replace("**", "").strip() or " "
                sanitized = _sanitize(text)
                if not sanitized:
                    sanitized = " "
                pdf.multi_cell(0, 6, txt=sanitized, align="L")

        # Flush any remaining table at the end
        if in_table:
            flush_table()

        return bytes(pdf.output())

    async def _convert_to_docx(self, markdown_text: str, title: str, show_watermark: bool = False, watermark_path: str | None = None) -> bytes:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import parse_xml
        from .export_templates import get_docx_watermark_anchor_xml

        logo_path = watermark_path or os.path.join(os.path.dirname(os.path.abspath(__file__)), "kreesalis.png")
        doc = Document()
        
        if show_watermark:
            # Add centered watermark to header
            for section in doc.sections:
                header = section.header
                # Ensure header has a paragraph
                if not header.paragraphs:
                    header.add_paragraph()
                header_para = header.paragraphs[0]
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                run = header_para.add_run()
                try:
                    # Larger, slightly more visible logo
                    picture = run.add_picture(logo_path, width=Inches(4))
                    
                    # Convert inline shape to floating shape behind text
                    inline = picture._inline
                    anchor_xml = get_docx_watermark_anchor_xml(inline.extent.cx, inline.extent.cy, inline.docPr.id, inline.docPr.name)
                    anchor = parse_xml(anchor_xml)
                    anchor.append(inline.graphic)
                    inline.getparent().replace(inline, anchor)
                except Exception as e:
                    self._logger.error(f"Failed to add picture to docx: {e}")

        doc.add_heading(title, 0)

        # Very basic markdown to docx parser
        lines = markdown_text.split("\n")
        for line_raw in lines:
            line: str = line_raw.strip()
            if not line:
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
